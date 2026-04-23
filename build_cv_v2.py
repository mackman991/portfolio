"""
Generate Peter_Mackin_CV_v2.docx  –  2-page target
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins (tight for 2-page fit) ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin   = Cm(1.4)
    section.right_margin  = Cm(1.4)

DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT = RGBColor(0x0F, 0x3D, 0x7A)
GREY   = RGBColor(0x44, 0x44, 0x44)

BODY   = 9       # pt  – body text
SMALL  = 8.5     # pt  – secondary / dates
LINE   = Pt(11)  # exact line spacing

def _base(p):
    """Apply tight single line spacing to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    spng = OxmlElement("w:spacing")
    spng.set(qn("w:line"),     "220")   # 220 twips ≈ 11pt
    spng.set(qn("w:lineRule"), "exact")
    pPr.append(spng)

def add_rule(p, color_hex="0F3D7A"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(1)
    _base(p)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size  = Pt(BODY)
    r.font.color.rgb = ACCENT
    r.font.name  = "Calibri"
    add_rule(p)
    return p

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before       = Pt(0)
    p.paragraph_format.space_after        = Pt(0)
    p.paragraph_format.left_indent        = Cm(0.38)
    p.paragraph_format.first_line_indent  = Cm(-0.38)
    _base(p)
    r = p.add_run(text)
    r.font.size  = Pt(BODY)
    r.font.color.rgb = GREY
    r.font.name  = "Calibri"
    return p

def para(text, size_pt=BODY, bold=False, color=None, space_after=1, space_before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    _base(p)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size  = Pt(size_pt)
    r.font.color.rgb = color if color else GREY
    r.font.name  = "Calibri"
    return p

def _add_right_tab(p, twips="9500"):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), twips)
    tabs.append(tab)
    pPr.append(tabs)

def job_header(company, title, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)
    _base(p)
    r1 = p.add_run(company + "  ")
    r1.bold = True; r1.font.size = Pt(BODY + 0.5); r1.font.color.rgb = DARK; r1.font.name = "Calibri"
    r2 = p.add_run("│  " + title)
    r2.font.size = Pt(BODY); r2.font.color.rgb = ACCENT; r2.font.name = "Calibri"
    p.add_run("\t")
    r3 = p.add_run(dates)
    r3.italic = True; r3.font.size = Pt(SMALL); r3.font.color.rgb = GREY; r3.font.name = "Calibri"
    _add_right_tab(p)
    return p

def project_header(title, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(0)
    _base(p)
    r1 = p.add_run(title + "  ")
    r1.bold = True; r1.font.size = Pt(BODY); r1.font.color.rgb = DARK; r1.font.name = "Calibri"
    r2 = p.add_run(tech)
    r2.italic = True; r2.font.size = Pt(SMALL); r2.font.color.rgb = ACCENT; r2.font.name = "Calibri"
    return p

# ═══════════════════════════════════════════════════════════════════════════════
#  NAME & CONTACT
# ═══════════════════════════════════════════════════════════════════════════════
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after  = Pt(2)
r = p_name.add_run("Peter Mackin")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK; r.font.name = "Calibri"

p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_after = Pt(1)
_base(p_contact)
rc = p_contact.add_run(
    "West Hampstead, London  •  pete.mackin12@gmail.com  •  07970 337464\n"
    "linkedin.com/in/peter-m-5450aa73  •  github.com/mackman991"
)
rc.font.size = Pt(SMALL); rc.font.color.rgb = GREY; rc.font.name = "Calibri"

# ═══════════════════════════════════════════════════════════════════════════════
#  PROFILE
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Profile")
para(
    "Investment professional combining seven years of multi-asset portfolio management, "
    "performance attribution and manager research (CFA Level II, IMC) with hands-on Python, "
    "SQL and machine-learning skills (LSE Data Analytics Career Accelerator, Distinction). "
    "Builds and deploys LLM-powered agent workflows using the Claude API to automate investment "
    "research, reporting and data pipelines; authored a quantitative earnings-signal study across "
    "433 events and four systematic strategy variants. Seeking roles where deep investment "
    "judgement and applied AI/data engineering create measurable edge.",
    space_after=1
)

# ═══════════════════════════════════════════════════════════════════════════════
#  CORE COMPETENCIES
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Core Competencies")

comp_rows = [
    ("AI & LLM Engineering",
     "Claude API  |  LLM agent design & orchestration  |  Agentic workflow automation  |  "
     "Prompt engineering  |  Claude Code  |  RAG  |  Financial data API integration"),
    ("Data & Analytics",
     "Python (pandas, scikit-learn, statsmodels, NLTK)  |  SQL  |  R  |  Tableau  |  Power BI  |  "
     "ML (regression, decision trees, k-means)  |  NLP & sentiment analysis  |  Predictive modelling"),
    ("Investments",
     "Portfolio Management  |  Performance Analysis & Attribution  |  Asset Allocation  |  "
     "Fund & Manager Selection  |  ESG Screening  |  Risk & Benchmarking  |  GIPS  |  "
     "Macro Research  |  Discretionary Trading (FX, Equities, FI, Commodities, Crypto)"),
    ("Professional",
     "Client Reporting  |  Committee Presentations  |  Strategic Advisory  |  "
     "Process Automation  |  Stakeholder Management"),
]
for label, content in comp_rows:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(0)
    _base(p)
    rb = p.add_run(label + ":  ")
    rb.bold = True; rb.font.size = Pt(BODY); rb.font.color.rgb = DARK; rb.font.name = "Calibri"
    rc2 = p.add_run(content)
    rc2.font.size = Pt(BODY); rc2.font.color.rgb = GREY; rc2.font.name = "Calibri"

# ═══════════════════════════════════════════════════════════════════════════════
#  SELECTED QUANTITATIVE & DATA PROJECTS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Selected Quantitative & Data Projects")

project_header(
    "Alpha Analytics  –  Earnings Event Study & Trading Strategies  (Capstone)",
    "Python  |  Event-study  |  Systematic strategy design  |  Streamlit  |  FMP API  |  SEC EDGAR"
)
for b in [
    "Built an end-to-end quantitative research pipeline across the S&P 100 (2,234 earnings events, "
    "2020–2025): daily OHLCV via yfinance, EPS surprises via FMP, SEC XBRL fundamentals and FRED macro "
    "series; engineered event windows, abnormal returns and regime-filter features for each print.",
    "Tested four systematic strategies; best result — Hold/Cut with MACD-bullish regime filter: "
    "+6.07% mean return per event, 81.0% hit-rate, t-stat 22.31. Under realistic portfolio sizing "
    "(3% per event, 30% gross cap): CAGR 16.6%, Sharpe 3.90, max drawdown −1.1%.",
    "Repackaged analysis into a modular Python library with Streamlit dashboard, regime-filter grid, "
    "position-sizing simulator and live monitoring module tracking upcoming events and rolling edge metrics.",
]:
    bullet(b)

project_header(
    "Customer Analytics & Predictive Modelling  –  Turtle Games (DA301)",
    "Python  |  R  |  scikit-learn  |  NLP (VADER)  |  k-means"
)
for b in [
    "Compared linear, multiple-linear and decision-tree regression on 2,000 customer records; "
    "the decision tree achieved R² = 0.96 predicting loyalty points from income and spending score, "
    "materially outperforming linear models that exhibited heteroscedasticity.",
    "Applied k-means clustering (elbow + silhouette) to identify five customer segments; "
    "used VADER sentiment analysis on product reviews to quantify polarity and surface marketing themes.",
]:
    bullet(b)

project_header(
    "NHS Capacity & Resource Utilisation (DA201)",
    "Python  |  Tableau  |  Diagnostic analysis  |  Scenario modelling"
)
for b in [
    "Analysed NHS appointment data across 42 ICBs and 106 sub-locations (2020–2022); the 1.2m daily "
    "guideline was breached on 175 of 334 days — identified seasonal peaks, regional disparities and "
    "the COVID-driven pivot to telephone consultations.",
    "Built scenario analysis on 3% vs 6% missed-appointment rates; delivered recommendations on "
    "data quality, resource allocation and workforce planning to senior NHS stakeholders.",
]:
    bullet(b)

# ═══════════════════════════════════════════════════════════════════════════════
#  WORK HISTORY
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Work History")

job_header("PMCL Consulting", "Senior Investment Associate", "Feb 2025 – Present")
for b in [
    "Designing and building an ecosystem of LLM-powered agents using the Claude API and Python to "
    "automate investment research, reporting and data workflows — including live financial data "
    "pipelines via REST APIs — freeing analyst capacity for higher-value advisory work.",
    "Lead portfolio analysis for charity and non-profit clients: performance analytics, asset "
    "allocation, strategic advice, and manager research and selection.",
    "Built ex-post and ex-ante forecasting models quantifying the impact of asset allocation and "
    "expenditure decisions on long-term risk and return outcomes.",
    "Present quarterly investment reports to client committees; guide trustees through manager "
    "selection, portfolio transitions and investment policy.",
]:
    bullet(b)

job_header("TTG Spinnaker", "Proprietary Trader", "Jan 2021 – Mar 2024")
for b in [
    "Executed a discretionary macro / event-driven strategy across FX, equities, fixed income, "
    "commodities and crypto, sizing positions around scheduled catalysts (central bank decisions, "
    "CPI releases, earnings) and technical setups.",
    "Managed risk across a multi-asset book: set position limits, monitored drawdown thresholds "
    "and adjusted exposure in response to changing volatility regimes.",
    "Conducted daily macro research, synthesising economic data, central bank communication and "
    "market positioning to generate trade ideas with defined entry, stop and target levels.",
]:
    bullet(b)

job_header("Stanhope Capital", "Investment Analyst / Associate", "Jun 2017 – Nov 2020")
for b in [
    "Calculated, analysed and modelled performance, risk and allocation data for consolidated "
    "multi-asset portfolios; ran multiple attribution models to isolate drivers of excess return "
    "across £4bn AUM — clients included British Heart Foundation, LSE and Duchy of Lancaster.",
    "Spearheaded the redesign and rebuild of the firm's performance database, cutting the monthly "
    "and quarterly reporting cycle and materially improving accuracy, scalability and data lineage.",
    "Delivered bespoke monthly and quarterly investment reports and presentations to charity and "
    "non-profit clients with detailed commentary on portfolios, holdings and the macro environment.",
    "Met with fund managers to evaluate strategy, process and fit against client objectives; assessed "
    "investments against ESG restrictions and ensured compliance with each client's IPS.",
]:
    bullet(b)

# ═══════════════════════════════════════════════════════════════════════════════
#  QUALIFICATIONS & CERTIFICATIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Qualifications & Certifications")
for q in [
    "CFA Level I & II (passed)  |  CFA Certificate in Performance Attribution  |  "
    "Investment Management Certificate (IMC)",
]:
    bullet(q)

# ═══════════════════════════════════════════════════════════════════════════════
#  EDUCATION
# ═══════════════════════════════════════════════════════════════════════════════
section_heading("Education")

edu_rows = [
    ("London School of Economics",  "Data Analytics Career Accelerator – Distinction", "2024"),
    ("Loughborough University",     "BSc Economics – 2:1",                             "2017"),
    ("The London Oratory School",   "A-Levels: AAB  |  GCSEs: 2A*, 6A, 2B",          ""),
]
for inst, detail, yr in edu_rows:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(0)
    _base(p)
    rb = p.add_run(inst + "  ")
    rb.bold = True; rb.font.size = Pt(BODY); rb.font.color.rgb = DARK; rb.font.name = "Calibri"
    rd = p.add_run(detail)
    rd.font.size = Pt(BODY); rd.font.color.rgb = GREY; rd.font.name = "Calibri"
    if yr:
        p.add_run("\t")
        ry = p.add_run(yr)
        ry.italic = True; ry.font.size = Pt(SMALL); ry.font.color.rgb = GREY; ry.font.name = "Calibri"
        _add_right_tab(p)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "c:/Users/petem/Claude/Projects/pm/Peter_Mackin_CV_v2.docx"
doc.save(out)
print(f"Saved: {out}")
